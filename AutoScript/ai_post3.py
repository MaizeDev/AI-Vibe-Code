import datetime
import argparse
import subprocess
import json
import os
import re
import hashlib
import requests
import openai
from pathlib import Path

# =========================
# 配置部分
# =========================
API_KEY = os.getenv("AI_POST_KEY", "sk-BIMlxWA1ksae6qASYOBDFlW1e4xByrALU9DHOHevCOyAeuyJ")
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "https://x666.me/v1")
GEMINI_BASE_URL = os.getenv("GEMINI_BASE_URL", "https://x666.me")

DEFAULT_MODEL = "gemini-2.5-flash"
DEFAULT_GROUP = "level3"

# =========================
# 核心功能：多格式文件读取 (增强版)
# =========================

def read_file_content(file_path: Path) -> str:
    """
    根据文件后缀名，智能读取不同格式的文件内容
    支持: .txt, .md, .docx, .pdf
    """
    suffix = file_path.suffix.lower()
    
    # 1. 处理 Word 文档 (.docx)
    if suffix == ".docx":
        try:
            import docx
        except ImportError:
            raise ImportError("读取 Word 需要安装库: pip install python-docx")
        
        print("📄 检测到 Word 文档，正在解析（含段落与表格）...")
        doc = docx.Document(file_path)
        full_text = []
        
        # 1.1 读取正文段落
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # 1.2 读取表格内容 (关键修复：防止表格内容读不到)
        for table in doc.tables:
            for row in table.rows:
                # 把每一行的单元格用 | 拼起来
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n\n".join(full_text)

    # 2. 处理 PDF 文档 (.pdf)
    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("读取 PDF 需要安装库: pip install pypdf")
            
        print("📄 检测到 PDF 文档，正在解析...")
        reader = PdfReader(file_path)
        full_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)
        return "\n\n".join(full_text)

    # 3. 默认处理纯文本 (.txt, .md 等)
    else:
        try:
            return file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            try:
                return file_path.read_text(encoding="gbk")
            except Exception:
                raise ValueError("❌ 无法识别的文件编码，请确保是 UTF-8 或 GBK")

# =========================
# 工具函数 (补回了丢失的函数)
# =========================

def extract_json_block(text: str) -> dict:
    text = text.strip()
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL).strip()
    if "```" in text:
        match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
        if match: text = match.group(1)
    if not text.startswith("{"):
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if match: text = match.group(0)
    try:
        text = text.replace("，", ",").replace("：", ":").replace("“", '"').replace("”", '"')
        return json.loads(text)
    except json.JSONDecodeError:
        return {}

def make_slug(title: str) -> str:
    slug = title.strip()
    slug = re.sub(r"[^\w\u4e00-\u9fa5-]", "-", slug)
    slug = re.sub(r"-+", "-", slug).strip("-")
    if not slug: slug = hashlib.md5(title.encode("utf-8")).hexdigest()[:8]
    return slug

def generate_front_matter(title, date, tags=None, draft=False):
    """
    生成 Hugo/Hexo 所需的 YAML 头信息
    """
    tags = tags or []
    lines = [
        "---",
        f'title: "{title}"',
        f"date: {date.isoformat()}",
        f"draft: {str(draft).lower()}",
    ]
    if tags:
        lines.append("tags:")
        for t in tags:
            lines.append(f"  - {t}")
    lines.append("---")
    return "\n".join(lines) + "\n\n"

# =========================
# AI 调用
# =========================

def call_llm_metadata(client, model, group, prompt):
    messages = [{"role": "user", "content": prompt}]
    try:
        if model.startswith("gemini-"):
            url = f"{GEMINI_BASE_URL}/v1beta/models/{model}:generateContent"
            headers = {"Content-Type": "application/json", "x-goog-api-key": API_KEY}
            payload = {"contents": [{"role": "user", "parts": [{"text": prompt}]}]}
            resp = requests.post(url, headers=headers, json=payload, timeout=30)
            if resp.status_code != 200: return "{}"
            return resp.json()["candidates"][0]["content"]["parts"][0]["text"]
        else:
            response = client.chat.completions.create(
                model=model, messages=messages, temperature=0.7, extra_body={"group": group}
            )
            if hasattr(response, 'choices'): return response.choices[0].message.content
            return response["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"⚠️ AI 接口调用出错: {e}")
        return "{}"

def generate_metadata(draft_text, client, model, group):
    print("🤖 正在分析内容并生成标题/标签...")
    
    # [调试信息] 让你知道到底读了多少字
    print(f"🔍 [调试] 成功提取文本长度: {len(draft_text)} 字符")
    if len(draft_text) < 5:
        print("⚠️ [警告] 提取内容为空！请检查文档是否只有图片。")

    # 只取前 1500 字做摘要分析
    prompt = (
        "阅读以下文章内容，提取一个精炼的中文标题和 3-5 个标签。\n"
        "严格只输出 JSON 格式：\n"
        "{\"title\": \"文章标题\", \"tags\": [\"标签1\", \"标签2\"]}\n\n"
        f"内容摘要：{draft_text[:1500]}"
    )
    
    content = call_llm_metadata(client, model, group, prompt)
    data = extract_json_block(content)
    
    title = data.get("title", "未命名文章")
    tags = data.get("tags", [])
    
    # 兜底：如果 AI 没生成标题，尝试用第一行文字
    if title == "未命名文章" and len(draft_text) > 5:
        print("⚠️ AI 未返回有效标题，尝试使用文档第一行...")
        lines = [line.strip() for line in draft_text.strip().split('\n') if line.strip()]
        if lines:
            first_line = lines[0]
            if len(first_line) < 50: 
                title = first_line

    print(f"✅ 标题：{title}")
    return title, tags

# =========================
# 主流程
# =========================

def main():
    parser = argparse.ArgumentParser(description="多格式博客发布工具")
    parser.add_argument("draft_file", help="草稿文件路径 (.txt, .md, .docx, .pdf)")
    parser.add_argument("--draft", action="store_true", help="标记为草稿")
    parser.add_argument("--output_dir", default="content/posts", help="输出目录")
    parser.add_argument("--model", default=DEFAULT_MODEL, help="模型名称")
    parser.add_argument("--group", default=DEFAULT_GROUP, help="API 分组")
    parser.add_argument("--no-git", action="store_true", help="跳过 Git 操作")
    args = parser.parse_args()

    client = openai.OpenAI(api_key=API_KEY, base_url=OPENAI_BASE_URL)
    draft_path = Path(args.draft_file)

    if not draft_path.exists():
        print(f"❌ 错误：找不到文件 {args.draft_file}")
        return

    # 1. 读取内容
    try:
        draft_text = read_file_content(draft_path)
        if not draft_text.strip():
            print("❌ 错误：文件内容为空或无法提取文本（请检查是否全是图片）")
            return
    except Exception as e:
        print(f"❌ 读取文件失败: {e}")
        return

    # 2. 生成元数据
    title, tags = generate_metadata(draft_text, client, args.model, args.group)
    
    # 3. 组合最终内容 (MD格式)
    now = datetime.datetime.now().astimezone()
    final_content = generate_front_matter(title, now, tags, args.draft) + draft_text

    # 4. 写入 Markdown 文件
    filename = f"{now.date()}-{make_slug(title)}.md"
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    filepath = output_dir / filename
    
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(final_content)
    
    print(f"💾 已将 [{draft_path.name}] 转换为 Markdown 并保存至：\n   -> {filepath}")

    # 5. Git 操作
    if not args.no_git:
        print("🚀 正在执行 Git 推送...")
        try:
            subprocess.run(["git", "add", str(filepath)], check=True, capture_output=True)
            subprocess.run(["git", "commit", "-m", f"Post: {title}"], check=True, capture_output=True)
            subprocess.run(["git", "push"], check=True, capture_output=True)
            print("✅ Git 推送成功！")
        except subprocess.CalledProcessError as e:
            print("❌ Git 操作失败！")
            print(f"错误详情：{e.stderr.decode('utf-8', errors='ignore')}")

if __name__ == "__main__":
    main()